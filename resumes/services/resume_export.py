"""Resume export services for the resumes app."""

from __future__ import annotations

import logging
from enum import StrEnum
from io import BytesIO
from pathlib import Path
from typing import Any

from django.conf import settings
from django.contrib.staticfiles import finders
from django.contrib.staticfiles.storage import staticfiles_storage
from django.template.loader import render_to_string

from resumes.models import Experience, Profile, Theme

logger = logging.getLogger(__name__)


class ExportFormat(StrEnum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


CONTENT_TYPES = {
    ExportFormat.PDF: "application/pdf",
    ExportFormat.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ExportFormat.TXT: "text/plain; charset=utf-8",
}


def build_resume_context(profile: Profile, theme: str | None = None) -> dict[str, Any]:
    selected_theme = theme if theme in Theme.values else profile.chosen_theme
    profile = (
        Profile.objects.filter(pk=profile.pk)
        .prefetch_related(
            "experiences__achievements",
            "education",
            "skills",
            "certifications",
        )
        .get()
    )
    return {
        "profile": profile,
        "theme": selected_theme,
        "template_name": f"resumes/themes/{selected_theme}.html",
        "experiences": profile.experiences.all(),
        "education": profile.education.all(),
        "skills": profile.skills.all(),
        "certifications": profile.certifications.all(),
    }


def export_resume(profile: Profile, theme: str, export_format: ExportFormat) -> bytes:
    context = build_resume_context(profile, theme)
    if export_format == ExportFormat.PDF:
        return export_pdf(context)
    if export_format == ExportFormat.DOCX:
        return export_docx(context)
    if export_format == ExportFormat.TXT:
        return export_txt(context).encode("utf-8")
    raise ValueError(f"Unsupported export format: {export_format}")


def resolve_pdf_stylesheet_uri() -> str | None:
    """Return a file:// URI to app.css for WeasyPrint (dev finder or collected STATIC_ROOT)."""

    css_path = finders.find("resumes/css/app.css")
    if css_path:
        return Path(css_path).resolve().as_uri()

    try:
        stored = staticfiles_storage.path("resumes/css/app.css")
    except AttributeError, NotImplementedError, ValueError:
        stored = None
    if stored:
        path = Path(stored)
        if path.is_file():
            return path.resolve().as_uri()

    collected = Path(settings.STATIC_ROOT) / "resumes" / "css" / "app.css"
    if collected.is_file():
        return collected.resolve().as_uri()

    return None


def weasyprint_runtime_available() -> bool:
    """Return True when WeasyPrint can render PDFs with native libraries installed."""

    try:
        from weasyprint import HTML

        HTML(string="<p>ok</p>", base_url=str(settings.BASE_DIR)).write_pdf()
    except OSError:
        return False
    return True


def export_pdf(context: dict[str, Any]) -> bytes:
    theme = context.get("theme", "unknown")
    try:
        from weasyprint import HTML

        pdf_context = {
            **context,
            "pdf_stylesheet_href": resolve_pdf_stylesheet_uri(),
        }
        html = render_to_string("resumes/export_pdf.html", pdf_context)
        base_url = Path(settings.BASE_DIR).resolve().as_uri()
        pdf = HTML(string=html, base_url=base_url).write_pdf()
        logger.info("pdf_export renderer=weasyprint theme=%s bytes=%s", theme, len(pdf))
        return pdf
    except OSError as exc:
        logger.warning(
            "pdf_export renderer=fallback theme=%s reason=%s",
            theme,
            exc,
        )
        return _fallback_pdf(context)


def pdf_uses_weasyprint(pdf_bytes: bytes) -> bool:
    """Heuristic: themed WeasyPrint PDFs include a Producer comment; fallback does not."""

    return b"WeasyPrint" in pdf_bytes


def _experience_heading(experience: Experience) -> str:
    return f"{experience.title} - {experience.company}".strip(" -")


def export_docx(context: dict[str, Any]) -> bytes:
    from docx import Document

    profile = context["profile"]
    document = Document()
    document.add_heading(profile.full_name or "Resume", level=0)
    contact = " | ".join(
        item
        for item in [
            profile.email,
            profile.phone,
            profile.location,
            profile.linkedin_url,
            profile.portfolio_url,
        ]
        if item
    )
    if contact:
        document.add_paragraph(contact)
    if profile.professional_summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(profile.professional_summary)

    if context["experiences"]:
        document.add_heading("Experience", level=1)
    for experience in context["experiences"]:
        heading = _experience_heading(experience)
        if heading:
            document.add_heading(heading, level=2)
        meta = " | ".join(item for item in [experience.location, experience.date_range()] if item)
        if meta:
            document.add_paragraph(meta)
        for achievement in experience.achievements.all():
            document.add_paragraph(achievement.text, style="List Bullet")

    if context["education"]:
        document.add_heading("Education", level=1)
    for education in context["education"]:
        document.add_paragraph(
            " | ".join(
                item
                for item in [
                    education.institution,
                    education.degree,
                    education.field,
                    education.date_range(),
                    f"GPA {education.gpa}" if education.gpa else "",
                ]
                if item
            )
        )
        if education.notes:
            document.add_paragraph(education.notes)

    if context["skills"]:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(skill.name for skill in context["skills"] if skill.name))

    if context["certifications"]:
        document.add_heading("Certifications", level=1)
    for certification in context["certifications"]:
        document.add_paragraph(
            " | ".join(
                item
                for item in [
                    certification.name,
                    certification.issuing_body,
                    certification.credential_id,
                ]
                if item
            )
        )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def export_txt(context: dict[str, Any]) -> str:
    profile = context["profile"]
    lines = [profile.full_name or "Resume"]
    contact = " | ".join(
        item
        for item in [profile.email, profile.phone, profile.location, profile.linkedin_url, profile.portfolio_url]
        if item
    )
    if contact:
        lines.extend([contact, ""])
    if profile.professional_summary:
        lines.extend(["SUMMARY", profile.professional_summary, ""])

    if context["experiences"]:
        lines.append("EXPERIENCE")
    for experience in context["experiences"]:
        heading = _experience_heading(experience)
        if heading:
            lines.append(heading)
        meta = " | ".join(item for item in [experience.location, experience.date_range()] if item)
        if meta:
            lines.append(meta)
        for achievement in experience.achievements.all():
            lines.append(f"- {achievement.text}")
        lines.append("")

    if context["education"]:
        lines.append("EDUCATION")
    for education in context["education"]:
        lines.append(
            " | ".join(
                item
                for item in [education.institution, education.degree, education.field, education.date_range()]
                if item
            )
        )
    if context["education"]:
        lines.append("")

    if context["skills"]:
        lines.extend(["SKILLS", ", ".join(skill.name for skill in context["skills"] if skill.name), ""])

    if context["certifications"]:
        lines.append("CERTIFICATIONS")
    for certification in context["certifications"]:
        lines.append(" | ".join(item for item in [certification.name, certification.issuing_body] if item))

    return "\n".join(lines).strip() + "\n"


def _fallback_pdf(context: dict[str, Any]) -> bytes:
    """Generate a simple valid PDF when local native WeasyPrint libs are absent."""

    text = export_txt(context)
    raw_lines = text.splitlines()
    if len(raw_lines) > 48 or any(len(line) > 100 for line in raw_lines):
        logger.warning("PDF fallback truncated resume content; install WeasyPrint native deps for full output.")
    lines = [_pdf_escape(line[:100]) for line in raw_lines[:48]]
    stream_lines = ["BT", "/F1 10 Tf", "72 740 Td"]
    for index, line in enumerate(lines):
        if index:
            stream_lines.append("0 -14 Td")
        stream_lines.append(f"({line}) Tj")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1", errors="ignore")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
    ]

    pdf = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{number} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            "trailer\n" f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n" "startxref\n" f"{xref_offset}\n" "%%EOF\n"
        ).encode("ascii")
    )
    return bytes(pdf)


def _pdf_escape(value: str) -> str:
    normalized = value.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\\\n", "\n")
    escaped: list[str] = []
    for char in normalized:
        if char == "\\":
            escaped.append("\\\\")
        elif char == "(":
            escaped.append("\\(")
        elif char == ")":
            escaped.append("\\)")
        elif char == "\n":
            escaped.append("\\n")
        else:
            escaped.append(char)
    return "".join(escaped)
